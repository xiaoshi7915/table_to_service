"""
Word 文档解析器（DOC/DOCX）
"""
from typing import Dict, Any, List
from pathlib import Path
from loguru import logger

from .base_parser import BaseParser

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx 未安装，Word 文档解析功能将不可用")


class DOCParser(BaseParser):
    """Word 文档解析器"""
    
    def __init__(self, file_path: str):
        super().__init__(file_path)
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安装，请安装 python-docx")
    
    def parse(self) -> Dict[str, Any]:
        """解析 Word 文档"""
        try:
            doc = DocxDocument(self.file_path)
            
            # 提取文本内容
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 提取表格内容
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        tables_text.append(row_text)
            
            # 合并所有文本
            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n表格内容:\n" + "\n".join(tables_text)
            
            # 提取元数据
            metadata = self.extract_metadata()
            
            return {
                "text": full_text,
                "metadata": metadata,
                "paragraphs": paragraphs,
                "tables_count": len(doc.tables),
            }
        except Exception as e:
            logger.error(f"Word 文档解析失败: {e}")
            raise
    
    def extract_metadata(self) -> Dict[str, Any]:
        """提取 Word 文档元数据"""
        try:
            doc = DocxDocument(self.file_path)
            core_props = doc.core_properties
            
            metadata = {
                "title": core_props.title,
                "author": core_props.author,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "comments": core_props.comments,
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "last_modified_by": core_props.last_modified_by,
                "revision": core_props.revision,
            }
            
            # 清理 None 值
            return {k: v for k, v in metadata.items() if v is not None}
        except Exception as e:
            logger.warning(f"提取 Word 文档元数据失败: {e}")
            return {}
    
    @classmethod
    def supports(cls, file_path: str) -> bool:
        """检查是否支持 Word 文档"""
        return Path(file_path).suffix.lower() in ['.doc', '.docx']

